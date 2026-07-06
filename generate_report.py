import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def read_file(filepath):
    with open(filepath, 'r') as file:
        return file.read()
# Paths to the Java files
base_path = "src/main/java/com/example/demo/"
test_path = "src/test/java/com/example/demo/"
pedido_java = read_file(os.path.join(base_path, "Pedido.java"))
pedido_repo = read_file(os.path.join(base_path, "PedidoRepository.java"))
pedido_service = read_file(os.path.join(base_path, "PedidoService.java"))
pedido_controller = read_file(os.path.join(base_path, "PedidoController.java"))
pedido_repo_test = read_file(os.path.join(test_path, "PedidoRepositoryTest.java"))
pedido_controller_test = read_file(os.path.join(test_path, "PedidoControllerTest.java"))
doc = docx.Document()
# ----------------- DATOS DE ENCABEZADO -----------------
# We can use a table to neatly display the header data
header = doc.add_heading('GUIA PARA LAS PRÁCTICAS DE LABORATORIO, TALLER O CAMPO', 0)
header.alignment = WD_ALIGN_PARAGRAPH.CENTER
table = doc.add_table(rows=10, cols=2)
table.style = 'Table Grid'
data = [
    ("Departamento", "DEPARTAMENTO DE CIENCIAS DE LA COMPUTACIÓN-SD"),
    ("Carrera", "INGENIERIA EN TECNOLOGIAS DE INFORMACIÓN"),
    ("Asignatura", "PROGRAMACIÓN AVANZADA"),
    ("Período lectivo", "202650"),
    ("Nivel", "7"),
    ("Docente", "JOHN JAVIER CRUZ GARZÓN"),
    ("NRC", "29547"),
    ("Práctica N°", "5"),
    ("Laboratorio", "LAB LA 103"),
    ("Tema", "Principios AAA y Test de Frameworks: pruebas unitarias de repositorios y controladores en Spring Boot")
]
for i, (key, value) in enumerate(data):
    table.cell(i, 0).text = key
    table.cell(i, 1).text = value
doc.add_paragraph()
# ----------------- 1. INTRODUCCIÓN -----------------
doc.add_heading('1. INTRODUCCIÓN', level=1)
intro_text = (
    "En el contexto del desarrollo de software con Spring Boot, las pruebas unitarias permiten "
    "validar de manera aislada el correcto funcionamiento de las distintas capas de la aplicación. "
    "El patrón AAA (Arrange, Act, Assert) es una convención fundamental que estructura cada prueba "
    "en tres fases claras: preparación del entorno y los datos (Arrange), ejecución de la lógica "
    "a probar (Act) y verificación de los resultados (Assert). Esto se evidencia de manera clara "
    "en los tests implementados en esta práctica.\n\n"
    "Spring Boot proporciona herramientas específicas conocidas como \"slice tests\" (pruebas por capas) "
    "para enfocarse en porciones específicas de la aplicación. La anotación @DataJpaTest se emplea "
    "para cargar únicamente el contexto relacionado con la persistencia de datos (repositorios y entidades), "
    "aislándolo de los controladores y servicios. Por otro lado, @WebMvcTest se usa para levantar "
    "solamente la capa web (controladores) y su infraestructura MVC, sin involucrar a la base de datos, "
    "lo cual permite realizar pruebas más rápidas y confiables."
)
doc.add_paragraph(intro_text)

# ----------------- 2. OBJETIVOS -----------------
doc.add_heading('2. OBJETIVOS', level=1)
doc.add_paragraph("Objetivo General:", style='List Bullet')
doc.add_paragraph("Aplicar los principios AAA (Arrange, Act, Assert) en la implementación de pruebas unitarias para repositorios y controladores en aplicaciones Spring Boot, utilizando las anotaciones @DataJpaTest y @WebMvcTest.")

doc.add_paragraph("Objetivos Específicos:", style='List Bullet')
doc.add_paragraph("1. Diseñar e implementar el modelo de dominio y su capa de persistencia utilizando Spring Data JPA.")
doc.add_paragraph("2. Desarrollar pruebas unitarias para la capa de persistencia mediante @DataJpaTest aislando el contexto de la base de datos.")
doc.add_paragraph("3. Construir la capa de servicio y exponer la funcionalidad a través de un controlador REST.")
doc.add_paragraph("4. Implementar pruebas unitarias para la capa web utilizando @WebMvcTest y @MockBean para simular el comportamiento de los servicios subyacentes.")

# ----------------- 3. DESARROLLO DE ACTIVIDADES -----------------
doc.add_heading('3. DESARROLLO DE ACTIVIDADES', level=1)

def add_code_block(document, code_string):
    p = document.add_paragraph(code_string)
    p.style.font.name = 'Courier New'
    p.style.font.size = Pt(9)

# Actividad 1
doc.add_heading('Actividad 1: Modelo de dominio', level=2)
doc.add_paragraph("Descripción de lo realizado: Se creó la entidad 'Pedido' y su correspondiente repositorio 'PedidoRepository' que extiende de JpaRepository, incluyendo un método de búsqueda por estado.")
doc.add_paragraph("Código de Pedido.java:")
add_code_block(doc, pedido_java)
doc.add_paragraph("Código de PedidoRepository.java:")
add_code_block(doc, pedido_repo)

# Actividad 2
doc.add_heading('Actividad 2: Prueba unitaria del repositorio (AAA)', level=2)
doc.add_paragraph("Descripción de lo realizado: Se implementó un test para el repositorio utilizando @DataJpaTest y aplicando el patrón AAA para verificar el método findByEstado.")
doc.add_paragraph("Explicación del patrón AAA en esta actividad:\n"
                  "- Arrange: Se guardan explícitamente en la base de datos en memoria (H2) 3 pedidos, 2 con estado PENDIENTE y 1 PAGADO, preparando el escenario.\n"
                  "- Act: Se invoca la función a probar, pedidoRepository.findByEstado(\"PENDIENTE\").\n"
                  "- Assert: Se verifica que la lista resultante contenga exactamente 2 elementos mediante las aserciones de AssertJ.")
doc.add_paragraph("Código de PedidoRepositoryTest.java:")
add_code_block(doc, pedido_repo_test)

# Actividad 3
doc.add_heading('Actividad 3: Capa de servicio y controlador', level=2)
doc.add_paragraph("Descripción de lo realizado: Se desarrolló 'PedidoService' para contener la lógica de negocio básica invocando al repositorio, y 'PedidoController' para exponer el servicio a través de un endpoint REST (@GetMapping) mediante un parámetro @RequestParam.")
doc.add_paragraph("Código de PedidoService.java:")
add_code_block(doc, pedido_service)
doc.add_paragraph("Código de PedidoController.java:")
add_code_block(doc, pedido_controller)

# Actividad 4
doc.add_heading('Actividad 4: Prueba unitaria del controlador (AAA)', level=2)
doc.add_paragraph("Descripción de lo realizado: Se probó el controlador aislando la capa de servicio con @MockBean y validando la interacción web mediante MockMvc y @WebMvcTest.")
doc.add_paragraph("Explicación del patrón AAA en esta actividad:\n"
                  "- Arrange: A través de given().willReturn(), configuramos el mock del servicio (pedidoService) para que retorne una lista quemada que contiene a 'Ana' sin consultar a la base de datos real.\n"
                  "- Act: Se ejecuta una petición falsa GET a '/api/pedidos' pasando el parámetro 'estado=PENDIENTE' a través del mockMvc.perform().\n"
                  "- Assert: Mediante los encadenamientos .andExpect() validamos que el código de estado HTTP sea isOk() (200) y que el JSON de respuesta contenga el cliente 'Ana' en el índice 0.")
doc.add_paragraph("Código de PedidoControllerTest.java:")
add_code_block(doc, pedido_controller_test)

# Actividad 5
doc.add_heading('Actividad 5: Checkpoint', level=2)
doc.add_paragraph("Diferencia entre @DataJpaTest y @WebMvcTest, y el uso de @MockBean:")
doc.add_paragraph("@DataJpaTest es una herramienta diseñada para probar exclusivamente la capa de persistencia; levanta los repositorios, las entidades y una base de datos en memoria sin cargar controladores ni servicios. Por el contrario, @WebMvcTest prueba exclusivamente la capa web; carga los controladores y la infraestructura de Spring MVC para simular peticiones HTTP pero no levanta servicios ni repositorios de base de datos. Finalmente, se utilizó @MockBean en la Actividad 4 debido a que, al estar en un entorno @WebMvcTest, el contenedor no posee instancias reales de la capa de servicio (PedidoService); el @MockBean permitió inyectar una simulación (mock) para garantizar que el controlador fuese probado en total aislamiento sin requerir la base de datos.")

# ----------------- 4. RESULTADOS OBTENIDOS -----------------
doc.add_heading('4. RESULTADOS OBTENIDOS', level=1)
doc.add_paragraph("[INSERTAR CAPTURA DE PANTALLA: ejecución de PedidoRepositoryTest en verde]")
doc.add_paragraph("[INSERTAR CAPTURA DE PANTALLA: ejecución de PedidoControllerTest en verde]")
doc.add_paragraph("Enlace al repositorio Git: [INSERTAR ENLACE AL REPOSITORIO]")

# ----------------- 5. CONCLUSIONES -----------------
doc.add_heading('5. CONCLUSIONES', level=1)
doc.add_paragraph("La estrategia de utilizar pruebas de tipo 'slice' en Spring Boot demostró un beneficio clave: permite focalizar el testing. Al aislar la capa de persistencia con @DataJpaTest, se detectan de forma rápida y concisa los posibles errores en consultas a la base de datos. Mientras tanto, aislar la capa web con @WebMvcTest junto con simulaciones mediante @MockBean evita depender de una base de datos lenta o preconfigurada, lo que asegura que las reglas de ruteo, respuestas JSON y mapeo de controladores se comporten correctamente de manera independiente y en tiempos de ejecución mucho más rápidos que las pruebas integrales de contexto completo.")

# ----------------- 6. RECOMENDACIONES -----------------
doc.add_heading('6. RECOMENDACIONES', level=1)
doc.add_paragraph("1. Mantener un claro distanciamiento entre los tests de integración total y los 'slice tests'. Si bien los 'slice tests' (como @DataJpaTest y @WebMvcTest) garantizan rapidez, es necesario poseer siempre al menos una prueba (e.g. @SpringBootTest) para validar el flujo completo ('end-to-end').")
doc.add_paragraph("2. Emplear bases de datos en memoria (como H2) estrictamente en los entornos de pruebas y siempre garantizar que la estructura de prueba recree los casos extremos (vacíos, nulos) para fortalecer los Assertions.")

# ----------------- 7. FIRMAS -----------------
doc.add_heading('7. FIRMAS', level=1)
signatures = doc.add_table(rows=3, cols=3)

# Add headers
signatures.cell(0, 0).text = "________________________________"
signatures.cell(0, 1).text = "________________________________"
signatures.cell(0, 2).text = "________________________________"

signatures.cell(1, 0).text = "Ing. John Javier Cruz, Mgtr."
signatures.cell(1, 1).text = "Ing. Kevin Chuquitarco, Mgtr."
signatures.cell(1, 2).text = "Ing. Javier Cevallos, Mgtr."

signatures.cell(2, 0).text = "DOCENTE"
signatures.cell(2, 1).text = "COORD. DE ÁREA DE CONOCIMIENTO"
signatures.cell(2, 2).text = "JEFE DE LABORATORIO"

doc.save('Informe_Laboratorio_5.docx')
print("Document generated successfully.")
